# a=[1,2,3] #list
# b={"hello","World"} #集合set
# print(a,type(a))
# print(b,type(b))

# int = 66
# char=chr(int) #字元:請先輸入10進制66(B)，再轉字元類型
# print(char*3,type(char)) #輸出B*3次

# c = 2**3 #2的3次方
# print(c,type(c))

#將py檔轉成exe，參考文獻:https://pypi.org/project/auto-py-to-exe/   開啟執行檔指令:"auto-py-to-exe"
#參考文獻:https://ithelp.ithome.com.tw/articles/10225127
import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import xlrd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

# 表格首行背景色設置
def tabBgColor(table,cols,colorStr):
    shading_list = locals()
    for i in range(cols):
        shading_list['shading_elm_'+str(i)] = parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'),bgColor = colorStr))
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_'+str(i)])


# 建立視窗
window = tk.Tk()
# 隱藏窗口
window.withdraw()

# 獲取对话框的宽度和高度
# 這些方法返回的是對話框所需的大小，而不是實際顯示的大小。
dialog_width = window.winfo_reqwidth()
dialog_height = window.winfo_reqheight()

# 獲取當下電腦螢幕的寬度和高度
screen_width = window.winfo_screenwidth()
screen_height = window.winfo_screenheight()


# 儲存檔案對話框
file_path = filedialog.asksaveasfilename(defaultextension='.docx', filetypes=[("Word Documents", "*.docx")], title='另存新檔')

if file_path:
    document = Document()

    document.add_heading('建立一個word檔', level=0)

    p = document.add_paragraph('大家好，我叫 ')
    #add_run():可以設定粗體或斜體等等特殊格式，他是屬於 paragraph 之下的方法，所以必須搭配 paragraph 物件使用。
    p.add_run('"邱冠為"').bold = True
    p.add_run('  今年25歲，是一名學習者  ')
    p.add_run('end~.').italic = True #italic:是否為斜體字

    document.add_heading('Heading, level 1',level=1) #標題的級別（level）
    document.add_paragraph('This is an intense quote.', style='Intense Quote') #預定義的內建樣式包括標題、引文、標準段落等

    document.add_paragraph('符號清單樣式的段落', style='List Bullet')
    document.add_paragraph('數字清單樣式的段落', style='List Number')

    headers = ['Qty', 'Id', 'Desc'] #欄位
    data = [[3, 101, 'Spam'],[7, 422, 'Eggs'],[4, 631, 'Spam, spam, eggs, and spam']] #資料列list，不可更換索引位置

    colsNum = 3
    colorStr = 'E6E6FA'
    #建立表格欄位與列
    table = document.add_table(rows=1, cols=colsNum, style='Light Grid')
    
    tabBgColor(table,colsNum,colorStr)

    header_cells = table.rows[0].cells

    # enumerate: 用於在 headers 列表的每個項目上生成索引和值的配對。變數 i 代表當前項目的索引，而變數 header 包含對應的值
    # Set the header text and font color
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        if len(cell.paragraphs) == 0:
            cell.add_paragraph()
        cell.paragraphs[0].add_run(header).font.color.rgb = RGBColor(139, 69, 19)
        # 設置單元格內容的水平和垂直對齊方式為居中
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #新增資料至表格裡
    for row in data:
        row_cells = table.add_row().cells 
        for i, cell_value in enumerate(row):
            cell = row_cells[i]
            if len(cell.paragraphs) == 0:
                cell.add_paragraph()
            cell.paragraphs[0].add_run(str(cell_value))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_page_break() #分頁符號
    
    # 儲存檔案
    document.save(file_path)
    
    print("Word 檔案建立成功！檔案已儲存於:", file_path)
else:
    print("未選擇儲存位置!")






